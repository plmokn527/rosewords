import os
from PIL import Image
from docx import Document
from docx.shared import Inches

# 创建一个新的Word文档
doc = Document()

# 获取qrcodes目录下所有的文件
path = 'qrcodes'
files = [f for f in os.listdir(path) if os.path.isfile(os.path.join(path, f))]

# 按照4行4列的形式将图片添加到Word文档中
for i in range(0, len(files), 16):  # 一共16张图片（4行4列）为一组
    table = doc.add_table(rows=4, cols=4)  # 添加一个4行4列的表格
    for j in range(i, min(i + 16, len(files))):  # 对于每一组
        row = j % 16 // 4
        col = j % 16 % 4
        cell = table.cell(row, col)  # 计算图片应该放在哪个单元格中
        paragraph = cell.add_paragraph()  # 在单元格中添加一个段落
        run = paragraph.add_run()  # 在段落中添加一个运行对象
        run.add_picture(os.path.join(path, files[j]), width=Inches(1.5))  # 将图片添加到运行对象中
    doc.add_page_break()  # 每16张图片（4行4列）后添加一个分页符

# 保存Word文档
doc.save('qrcodes.docx')
